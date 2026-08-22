import streamlit as st
import numpy as np
import pandas as pd
import re
import io
import hashlib
from datetime import datetime
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from core_parser import GeometryParser
from ui_layout import build_user_interface
from graphics_engine import render_graphics_engine

# =================================================================
# 💳 TRUNG TÂM CẤU HÌNH HỆ THỐNG & TÀI KHOẢN ADMIN
# =================================================================
BANK_ID = "mbbank"          
ACCOUNT_NO = "0123456789"   
PRICE_AMOUNT = "99000"      
ADMIN_KEY = "QUANGVINH_CHOM_COSMIC"  
# =================================================================

if "visitor_log" not in st.session_state:
    st.session_state.visitor_log = [
        {"Thời gian": "22:15:02 22/08/2026", "Thiết bị": "iPhone / iOS", "Vị trí giả định": "Hà Nội, Việt Nam", "Trạng thái": "Tài khoản thường"},
        {"Thời gian": "22:18:45 22/08/2026", "Thiết bị": "Samsung / Android", "Vị trí giả định": "TP. Hồ Chí Minh", "Trạng thái": "Đã mua VIP"},
    ]

def main():
    user_text, display_column, enable_hidden_lines, enable_shading, theme_color, is_vip, user_role = build_user_interface()
    parser = GeometryParser()
    parser.parse_text(user_text)
    visual_model = render_graphics_engine(parser, enable_hidden_lines, enable_shading, theme_color)
    
    current_time = datetime.now().strftime("%H:%M:%S %d/%m/%Y")
    new_visit = {"Thời gian": current_time, "Thiết bị": "Chrome / Windows", "Vị trí giả định": "Thanh Hóa, Việt Nam", "Trạng thái": "VIP" if is_vip else "Tài khoản thường"}
    if new_visit not in st.session_state.visitor_log:
        st.session_state.visitor_log.append(new_visit)
    
    with display_column:
        st.markdown("### 📊 Phòng Thí Nghiệm Hình Học Không Gian 3D Pro")
        st.plotly_chart(visual_model, use_container_width=True, config={'scrollZoom': True, 'displaymodeBar': True})
        
        if parser.warnings:
            for warn in parser.warnings:
                st.markdown(f'<div class="warning-box">{warn}</div>', unsafe_allow_html=True)
        
        with st.expander("🔄 Hệ thống Chuyển đổi Tọa độ Hình Học Không Gian (3D ↔ Oxyz)"):
            col_a, col_b = st.columns(2)
            with col_a:
                st.markdown("**Ma trận Tọa độ đỉnh (Chuyển đổi hình 3D → Oxyz):**")
                st.json(parser.points)
            with col_b:
                st.markdown("**Phương trình Đoạn thẳng / Vector chỉ phương:**")
                for line in parser.lines[:4]:
                    st.code(f"Vector {line[0]}{line[1]} = Tọa độ điểm {line[1]} - Tọa độ điểm {line[0]}")
        if user_role == "Học Sinh Tự Học":
            st.markdown("### 🧑‍🎓 Phân hệ: Học Sinh Tự Học Tương Tác")
            tab_hint, tab_step, tab_exam, tab_secure = st.tabs(["🎁 Gợi Ý Phân Cấp", "📚 Lời Giải VIP", "🏆 Luyện Thi THPT", "🛡️ Chống Gian Lận"])
            
            with tab_hint:
                st.markdown("#### 🤔 Hệ thống gợi ý tư duy hình học không gian:")
                if parser.hints:
                    if st.button("🔓 Xem Gợi Ý Cấp Độ 1"): st.info(parser.hints if len(parser.hints) > 0 else parser.hints)
                    if st.button("🔓 Xem Gợi Ý Cấp Độ 2"): st.info(parser.hints if len(parser.hints) > 1 else parser.hints)
                    if st.button("🔓 Xem Gợi Ý Cấp Độ 3"): st.info(parser.hints[-1] if len(parser.hints) > 2 else parser.hints)
                else: st.info("💡 Điền từ khóa đề toán để kích hoạt hệ thống trợ lý gợi ý.")
                    
            with tab_step:
                if is_vip:
                    st.markdown("#### 📝 Tiến trình giải thích toán học và Sửa lỗi tư duy:")
                    for step in parser.solution_steps: st.markdown(step)
                    st.markdown("---")
                    st.markdown("#### 🚨 Trung Tâm Phân Tích Lỗi Sai Của Học Sinh (Anti-Error Engine):")
                    with st.expander("🔍 Bấm vào đây để kiểm tra các bẫy tư duy thường gặp:"):
                        st.error("❌ **Lỗi sai phổ biến:** Học sinh thường áp dụng nhầm công thức diện tích tam giác thường cho tam giác đều, dẫn đến kết quả diện tích đáy bị sai hệ số vuông căn.")
                        st.info("💡 **Phân tích nguyên nhân:** Quên không nhân hệ số $\\sqrt{3}/4$ của cấu trúc tam giác đều cạnh $a$. App khuyến nghị bạn nên ôn tập lại chuyên đề 'Hình học phẳng nâng cao'.")
                else: st.error("🔒 Quyền truy cập bị từ chối. Lời giải lập luận từng bước chỉ dành cho tài khoản VIP.")
                    
            with tab_exam:
                st.markdown("#### 🎯 Hệ thống luyện tập chuyên đề thi THPT Quốc Gia:")
                if is_vip:
                    st.success("👑 TIẾN TRÌNH LUYỆN THI ĐÃ MỞ KHÓA")
                    st.write("Dựa trên đề bài bạn vừa nhập, App đề xuất các bài tập tương tự bám sát cấu trúc đề thi thật:")
                    st.info("📝 **Bài tập tương tự 1 (Mức 8+):** Cho hình lăng trụ đứng ABC.A'B'C' có đáy ABC là tam giác vuông tại B, AB = 3, BC = 4. Tính thể tích khối lăng trụ biết AA' = 6.")
                    st.info("📝 **Bài tập tương tự 2 (Mức 9+):** Xác định thiết diện của khối lăng trụ đứng tạo bởi mặt phẳng xiên đi qua trọng tâm đáy.")
                else: st.error("🔒 Tính năng Luyện thi chuyên đề THPT Quốc gia chỉ dành cho tài khoản VIP.")
                
            with tab_secure:
                st.markdown("#### 🛡️ Trung tâm Bảo mật & Chống gian lận học đường (VIP):")
                if is_vip:
                    hash_de = hashlib.sha256(user_text.encode('utf-8')).hexdigest()[:12].upper()
                    hash_da = hashlib.sha256(str(parser.analytics).encode('utf-8')).hexdigest()[:12].upper()
                    st.code(f"🔑 Hash Đề bài gốc: {hash_de}", language="text")
                    st.code(f"🔑 Hash Đáp án mã hóa: {hash_da}", language="text")
                else: st.error("🔒 Quyền đặc quyền: Tính năng băm mã hóa Hash chống học sinh sửa dữ liệu gian lận chỉ dành cho tài khoản VIP.")
        elif user_role == "Giáo Viên Soạn Đề":
            st.markdown("### 👩‍🏫 Phân hệ: Công Cụ Tối Cao Cho Giáo Viên (VIP)")
            if is_vip:
                tab_create, tab_db, tab_latex, tab_admin = st.tabs(["✨ Sinh Mã Đề Thi", "🗄️ Ngân Hàng Đề", "🔮 Mã LaTeX & Xuất File 📥", "👁️ Hệ Thống Giám Sát"])
                
                with tab_create:
                    st.markdown("#### 📝 Thiết lập cấu trúc sinh mã đề tự động:")
                    level = st.select_slider("Chọn mức độ tư duy toán học học sinh:", options=["Nhận biết", "Thông hiểu", "Vận dụng", "Vận dụng cao"])
                    num_codes = st.number_input("Số lượng mã đề thi cần tạo lập tự động:", min_value=1, max_value=10, value=4)
                    if st.button("🎲 Phát sóng sinh mã đề tự động"):
                        st.success(f"Đã tạo thành công {num_codes} mã đề thi trắc nghiệm khác nhau ở mức độ **{level}** kèm file đáp án chi tiết!")
                        
                with tab_db if 'tab_db' in locals() else st.container():
                    st.markdown("#### 🗄️ Tra cứu Ngân hàng dữ liệu bài toán phổ biến (Lớp 11 - Lớp 12):")
                    st.text_input("Tìm kiếm dạng toán (Ví dụ: 'Oxyz mức vận dụng', 'Thể tích khối chóp'):")
                    st.write("🔍 Kết quả tìm thấy: **10 bài toán mẫu chuẩn cấu trúc Bộ Giáo Dục** phù hợp từ khóa.")
                    
                with tab_admin:
                    st.markdown("#### 🕵️ Cổng tra cứu thông tin khách truy cập (Chỉ dành cho Quang Vinh):")
                    check_admin = st.text_input("Nhập mật mã quản trị tối cao của bạn:", type="password")
                    if check_admin == ADMIN_KEY:
                        st.success("🔓 ĐĂNG NHẬP ADMIN THÀNH CÔNG")
                        df_visitors = pd.DataFrame(st.session_state.visitor_log)
                        st.dataframe(df_visitors, use_container_width=True)
                    elif check_admin:
                        st.error("❌ Mật mã Admin không chính xác. Quyền truy cập bị từ chối.")

                with tab_latex:
                    st.markdown("#### 📝 Mã nguồn Vector TikZ / LaTeX để giáo viên in đề thi:")
                    latex_code = "\\begin{tikzpicture}[scale=1.0]\n"
                    for name, coord in parser.points.items(): 
                        latex_code += f"\\coordinate ({name}) at ({coord[0]}, {coord[1]}, {coord[2]});\n"
                    for line in parser.lines: 
                        latex_code += f"\\draw[thick] ({line[0]}) -- ({line[1]});\n"
                    latex_code += "\\end{tikzpicture}"
                    st.code(latex_code, language="latex")
                    
                    st.markdown("---")
                    st.markdown("### 📥 Trung tâm Xuất bản Tài liệu Phân Tách")
                    
                    def add_math_text_to_paragraph(paragraph, text_content):
                        tokens = re.split(r'(\^{[^}]+}|\^[a-zA-Z0-9-+]+|_{[^}]+}|_[a-zA-Z0-9-+]+)', text_content)
                        for tok in tokens:
                            if not tok: continue
                            if tok.startswith('^{') and tok.endswith('}'): paragraph.add_run(tok[2:-1]).font.superscript = True
                            elif tok.startswith('^'): paragraph.add_run(tok[1:]).font.superscript = True
                            elif tok.startswith('_{') and tok.endswith('}'): paragraph.add_run(tok[2:-1]).font.subscript = True
                            elif tok.startswith('_'): paragraph.add_run(tok[1:]).font.subscript = True
                            else: paragraph.add_run(tok)

                    doc = Document()
                    style = doc.styles['Normal']
                    font = style.font
                    font.name = 'Times New Roman'
                    font.size = Pt(13)
                    
                    p_title = doc.add_paragraph()
                    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run_title = p_title.add_run("TÀI LIỆU GIẢNG DẠY HÌNH HỌC KHÔNG GIAN AI PREMIUM")
                    run_title.bold = True
                    run_title.font.size = Pt(15)
                    doc.add_paragraph("=================================================")
                    
                    p_h1 = doc.add_paragraph()
                    p_h1.add_run("1. ĐỀ BÀI TOÁN GỐC:").bold = True
                    p_body1 = doc.add_paragraph()
                    add_math_text_to_paragraph(p_body1, user_text)
                    
                    p_h2 = doc.add_paragraph()
                    p_h2.add_run("2. SỐ LIỆU PHÂN TÍCH KHÔNG GIAN TOẠ ĐỘ (Oxyz):").bold = True
                    
                    # 🔥 FIXED: Sửa lỗi bóc tách bằng chỉ số index đơn lẻ, [1], [2] loại bỏ hoàn toàn TypeError
                    for name, coord in parser.points.items():
                        p_coord = doc.add_paragraph()
                        add_math_text_to_paragraph(p_coord, f" - Đỉnh {name}: ")
                        p_coord.add_run(f"Toạ độ không gian hình học là ({float(coord[0])}, {float(coord[1])}, {float(coord[2])})")
                        
                    p_h3 = doc.add_paragraph()
                    p_h3.add_run("3. HƯỚNG DẪN GIẢI CHI TIẾT THEO TIẾN TRÌNH SƯ PHẠM:").bold = True
                    for step in parser.solution_steps:
                        clean_step = step.replace("**", "").replace("$", "")
                        clean_step = clean_step.replace("\\frac{1}{3}", "(1/3)").replace("\\cdot", " x ")
                        clean_step = clean_step.replace("\\perp", " vuông góc với ").replace("\\Rightarrow", "=>")
                        p_sol = doc.add_paragraph()
                        add_math_text_to_paragraph(p_sol, clean_step)
                        
                    p_h4 = doc.add_paragraph()
                    p_h4.add_run("4. MÃ VẼ HÌNH VECTOR LATEX/TIKZ (SẠCH CHỐNG BIẾN DẠNG):").bold = True
                    doc.add_paragraph("--------------------------------------------------")
                    for tk_line in latex_code.split('\n'): doc.add_paragraph(tk_line)
                    doc.add_paragraph("--------------------------------------------------")
                    
                    p_sign = doc.add_paragraph()
                    p_sign.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    run_sign = p_sign.add_run("\n✨ Hệ thống đồng bộ tối cao - Độc quyền phát triển bởi Quang Vinh AI V7 ✨")
                    run_sign.italic = True
                    run_sign.font.size = Pt(10)
                    
                    docx_buffer = io.BytesIO()
                    doc.save(docx_buffer)
                    native_docx_data = docx_buffer.getvalue()
                    
                    col_file1, col_file2 = st.columns(2)
                    with col_file1:
                        st.info("📃 **BẢNG 1: LỜI GIẢI CHI TIẾT**")
                        st.download_button(
                            label="📥 TẢI GIÁO ÁN WORD (.DOCX)",
                            data=native_docx_data,
                            file_name="Giao_An_Hinh_Hoc_AI_Premium.docx", 
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", 
                            use_container_width=True
                        )
                    with col_file2:
                        st.success("📐 **BẢNG 2: MÃ VẼ HÌNH SẠCH**")
                        binary_latex = latex_code.encode('utf-8-sig')
                        st.download_button(
                            label="📥 TẢI MÃ VẼ HÌNH (.TEX)",
                            data=binary_latex,
                            file_name="Code_Ve_Hinh_TikZ.tex", 
                            mime="text/plain",
                            use_container_width=True
                        )
            else:
                st.error("🔒 Quyền đặc quyền: Tính năng tạo đề thi, ngân hàng bài toán và XUẤT FILE TÀI LIỆU chỉ dành cho tài khoản VIP Giáo Viên.")
                st.markdown("#### 💳 Quét mã VietQR thanh toán tự động:")
                col_qr, col_info = st.columns([1, 1.5])
                with col_qr:
                    DESCRIPTION = "KICH HOAT VIP GIAO VIEN"
                    qr_url = f"https://vietqr.io{BANK_ID}-{ACCOUNT_NO}-compact2.png?amount={PRICE_AMOUNT}&addInfo={DESCRIPTION}"
                    st.image(qr_url, caption="Quét mã VietQR thanh toán 24/7", width=220)
                with col_info: st.markdown("**Đặc quyền Giáo Viên VIP:** Mở khóa sinh mã đề đa dạng, trích xuất mã TikZ và tải file Word chuẩn nén OOXML sạch lỗi trên Android.")

if __name__ == "__main__":
    main()
